"""
Document parsing utilities for comprehensive file type support.

This module provides robust, production-ready parsers for 50+ file formats
including documents, spreadsheets, code, notebooks, email, archives, and more.

Features:
- 50+ supported file extensions across 10 categories
- Graceful fallbacks when optional dependencies unavailable
- Memory-efficient processing for large files
- Comprehensive error handling with detailed messages
- Metadata extraction (page counts, row counts, etc.)
- Encoding detection for text files
- Archive inspection and text extraction

Supported Categories:
- Text: .txt, .md, .rst, .tex, .org, .adoc, etc.
- Code: .py, .js, .ts, .java, .go, .rs, etc.
- Documents: .pdf, .docx, .doc, .odt, .rtf
- Spreadsheets: .xlsx, .xls, .ods, .csv, .tsv
- Web: .html, .htm, .xml, .rss, .atom
- Notebooks: .ipynb, .rmd, .qmd
- Email: .eml, .msg, .mbox
- Archives: .zip, .tar, .tar.gz, .7z
- Data: .json, .jsonl, .yaml, .toml
- Config: .conf, .env, .ini

Usage:
    from shared.utils import parse_file, is_supported_file

    # Parse any supported file
    result = parse_file("document.pdf")
    print(result['content'])
    print(result['metadata'])

    # Check if file is supported
    if is_supported_file("example.docx"):
        result = parse_file("example.docx")

Author: Luke Steuber
"""

import os
import re
import json
import csv
import logging
from io import StringIO
from pathlib import Path
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional, Set, Union

# Optional imports with runtime checks
try:
    from pdfminer.high_level import extract_text as pdf_extract_text
    from pdfminer.pdfparser import PDFSyntaxError
    PDF_AVAILABLE = True
except ImportError:
    pdf_extract_text = None
    PDFSyntaxError = None
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    Document = None
    DOCX_AVAILABLE = False

try:
    from openpyxl import load_workbook
    EXCEL_AVAILABLE = True
except ImportError:
    load_workbook = None
    EXCEL_AVAILABLE = False

try:
    import xlrd
    XLS_AVAILABLE = True
except ImportError:
    xlrd = None
    XLS_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    HTML_AVAILABLE = True
except ImportError:
    BeautifulSoup = None
    HTML_AVAILABLE = False

try:
    import pypandoc
    PANDOC_AVAILABLE = True
except ImportError:
    pypandoc = None
    PANDOC_AVAILABLE = False

try:
    from email.parser import BytesParser
    from email.policy import default as email_default_policy
    EMAIL_AVAILABLE = True
except ImportError:
    BytesParser = None
    email_default_policy = None
    EMAIL_AVAILABLE = False

try:
    import zipfile
    import tarfile
    ARCHIVE_AVAILABLE = True
except ImportError:
    zipfile = None
    tarfile = None
    ARCHIVE_AVAILABLE = False

logger = logging.getLogger(__name__)


# ============================================================================
# Data Classes
# ============================================================================

@dataclass
class ParseResult:
    """
    Result from file parsing operation.

    Attributes:
        content: Extracted text content
        metadata: File metadata (size, type, pages, etc.)
        success: Whether parsing succeeded
        error: Error message if parsing failed
    """
    content: str
    metadata: Dict[str, Any]
    success: bool
    error: Optional[str] = None


# ============================================================================
# File Type Constants
# ============================================================================

# Supported file extensions by category
TEXT_EXTENSIONS: Set[str] = {
    '.txt', '.md', '.rst', '.tex', '.org', '.adoc', '.wiki',
    '.markdown', '.mdown', '.mkd', '.text', '.asc',
    '.textile', '.mediawiki', '.creole', '.bbcode'
}

CODE_EXTENSIONS: Set[str] = {
    '.py', '.js', '.ts', '.jsx', '.tsx', '.html', '.htm', '.css',
    '.json', '.xml', '.yaml', '.yml', '.toml', '.ini', '.cfg',
    '.sql', '.sh', '.bash', '.zsh', '.fish', '.ps1', '.bat',
    '.php', '.rb', '.go', '.rs', '.cpp', '.c', '.h', '.hpp',
    '.java', '.scala', '.kt', '.swift', '.dart', '.r', '.m',
    '.pl', '.lua', '.vim', '.dockerfile', '.makefile', '.gradle',
    '.cmake', '.scss', '.sass', '.less', '.styl', '.vue', '.svelte',
    '.astro', '.mjs', '.cjs', '.coffee', '.litcoffee'
}

WEB_EXTENSIONS: Set[str] = {
    '.html', '.htm', '.xhtml', '.xml', '.rss', '.atom', '.svg',
    '.wml', '.xsl', '.xslt', '.jsp', '.asp', '.aspx'
}

DATA_EXTENSIONS: Set[str] = {
    '.csv', '.tsv', '.jsonl', '.ndjson', '.log', '.parquet',
    '.arrow', '.feather', '.pickle', '.pkl', '.hdf5', '.h5'
}

DOCUMENT_EXTENSIONS: Set[str] = {
    '.pdf', '.docx', '.doc', '.odt', '.pages', '.rtf'
}

SPREADSHEET_EXTENSIONS: Set[str] = {
    '.xlsx', '.xls', '.xlsm', '.ods', '.numbers', '.gnumeric'
}

NOTEBOOK_EXTENSIONS: Set[str] = {
    '.ipynb', '.rmd', '.qmd', '.rmarkdown'
}

EMAIL_EXTENSIONS: Set[str] = {
    '.eml', '.msg', '.mbox', '.maildir'
}

ARCHIVE_EXTENSIONS: Set[str] = {
    '.zip', '.tar', '.tar.gz', '.tgz', '.tar.bz2', '.tbz2',
    '.tar.xz', '.txz', '.rar', '.7z'
}

CONFIG_EXTENSIONS: Set[str] = {
    '.conf', '.config', '.properties', '.env', '.editorconfig',
    '.gitignore', '.dockerignore', '.eslintrc', '.prettierrc',
    '.babelrc', '.tsconfig', '.package', '.lock'
}


# ============================================================================
# File Parser Class
# ============================================================================

class FileParser:
    """
    High-performance file parser with support for 50+ formats.

    Example:
        >>> parser = FileParser()
        >>> result = parser.parse_file("document.pdf")
        >>> if result.success:
        ...     print(result.content)
        ...     print(f"Pages: {result.metadata.get('pages', 'N/A')}")
    """

    @classmethod
    def get_supported_extensions(cls) -> Set[str]:
        """Get all supported file extensions."""
        return (
            TEXT_EXTENSIONS | CODE_EXTENSIONS | WEB_EXTENSIONS |
            DATA_EXTENSIONS | DOCUMENT_EXTENSIONS | SPREADSHEET_EXTENSIONS |
            NOTEBOOK_EXTENSIONS | EMAIL_EXTENSIONS | ARCHIVE_EXTENSIONS |
            CONFIG_EXTENSIONS
        )

    @classmethod
    def is_supported(cls, file_path: Union[str, Path]) -> bool:
        """Check if file type is supported."""
        ext = Path(file_path).suffix.lower()
        return ext in cls.get_supported_extensions()

    @classmethod
    def get_file_type(cls, file_path: Union[str, Path]) -> str:
        """Determine file type category."""
        ext = Path(file_path).suffix.lower()

        if ext in TEXT_EXTENSIONS:
            return "text"
        elif ext in CODE_EXTENSIONS:
            return "code"
        elif ext in WEB_EXTENSIONS:
            return "web"
        elif ext in DATA_EXTENSIONS:
            return "data"
        elif ext in DOCUMENT_EXTENSIONS:
            return "document"
        elif ext in SPREADSHEET_EXTENSIONS:
            return "spreadsheet"
        elif ext in NOTEBOOK_EXTENSIONS:
            return "notebook"
        elif ext in EMAIL_EXTENSIONS:
            return "email"
        elif ext in ARCHIVE_EXTENSIONS:
            return "archive"
        elif ext in CONFIG_EXTENSIONS:
            return "config"
        else:
            return "unknown"

    def parse_file(self, file_path: Union[str, Path]) -> ParseResult:
        """
        Parse file and extract text content with metadata.

        Args:
            file_path: Path to file to parse

        Returns:
            ParseResult with content, metadata, and status

        Example:
            >>> parser = FileParser()
            >>> result = parser.parse_file("data.csv")
            >>> print(f"Rows: {result.metadata['rows']}")
        """
        file_path = Path(file_path)

        # Check file exists
        if not file_path.exists():
            return ParseResult(
                content='',
                metadata={'file_path': str(file_path)},
                success=False,
                error=f"File not found: {file_path}"
            )

        # Initialize metadata
        metadata = {
            'file_path': str(file_path),
            'file_name': file_path.name,
            'file_size': file_path.stat().st_size,
            'file_type': self.get_file_type(file_path),
            'extension': file_path.suffix.lower(),
            'encoding': 'utf-8'
        }

        try:
            # Parse based on file extension
            ext = file_path.suffix.lower()

            if ext == '.pdf':
                content, extra_meta = self._parse_pdf(file_path)
            elif ext in {'.docx', '.doc'}:
                content, extra_meta = self._parse_docx(file_path)
            elif ext in {'.xlsx', '.xls', '.xlsm'}:
                content, extra_meta = self._parse_excel(file_path)
            elif ext == '.csv':
                content, extra_meta = self._parse_csv(file_path)
            elif ext == '.tsv':
                content, extra_meta = self._parse_tsv(file_path)
            elif ext in {'.json', '.jsonl', '.ndjson'}:
                content, extra_meta = self._parse_json(file_path)
            elif ext == '.ipynb':
                content, extra_meta = self._parse_notebook(file_path)
            elif ext in {'.html', '.htm', '.xhtml'}:
                content, extra_meta = self._parse_html(file_path)
            elif ext in {'.xml', '.rss', '.atom', '.svg'}:
                content, extra_meta = self._parse_xml(file_path)
            elif ext in {'.eml', '.msg'}:
                content, extra_meta = self._parse_email(file_path)
            elif ext == '.rtf':
                content, extra_meta = self._parse_rtf(file_path)
            elif ext in ARCHIVE_EXTENSIONS:
                content, extra_meta = self._parse_archive(file_path)
            elif ext in TEXT_EXTENSIONS | CODE_EXTENSIONS | CONFIG_EXTENSIONS:
                content, extra_meta = self._parse_text(file_path)
            else:
                # Try as text file
                content, extra_meta = self._parse_text(file_path)

            # Merge metadata
            metadata.update(extra_meta)

            return ParseResult(
                content=content,
                metadata=metadata,
                success=True
            )

        except Exception as e:
            logger.error(f"Error parsing {file_path}: {e}")
            return ParseResult(
                content='',
                metadata=metadata,
                success=False,
                error=str(e)
            )

    # ========================================================================
    # Format-Specific Parsers
    # ========================================================================

    def _parse_pdf(self, file_path: Path) -> tuple[str, dict]:
        """Parse PDF files using pdfminer."""
        if not PDF_AVAILABLE:
            raise ImportError("PDF parsing requires pdfminer.six: pip install pdfminer.six")

        try:
            content = pdf_extract_text(
                str(file_path),
                maxpages=0,
                caching=True,
                codec='utf-8'
            )

            content = self._clean_text(content)

            if not content or len(content.strip()) < 10:
                raise ValueError("PDF appears empty or contains only images")

            return content, {'pages': content.count('\f') + 1}

        except PDFSyntaxError:
            raise ValueError("PDF file corrupted or invalid")
        except Exception as e:
            if "password" in str(e).lower():
                raise ValueError("PDF is password protected")
            elif "encrypted" in str(e).lower():
                raise ValueError("PDF is encrypted")
            raise

    def _parse_docx(self, file_path: Path) -> tuple[str, dict]:
        """Parse DOCX files using python-docx."""
        if not DOCX_AVAILABLE:
            raise ImportError("DOCX parsing requires python-docx: pip install python-docx")

        doc = Document(str(file_path))

        # Extract paragraphs
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        # Extract tables
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    tables_text.append(row_text)

        # Combine
        content_parts = paragraphs
        if tables_text:
            content_parts.extend(['', '--- Tables ---'] + tables_text)

        content = '\n'.join(content_parts)

        return content, {
            'paragraphs': len(paragraphs),
            'tables': len(doc.tables)
        }

    def _parse_excel(self, file_path: Path) -> tuple[str, dict]:
        """Parse Excel files (.xlsx, .xls)."""
        ext = file_path.suffix.lower()

        if ext in {'.xlsx', '.xlsm'}:
            if not EXCEL_AVAILABLE:
                raise ImportError("XLSX parsing requires openpyxl: pip install openpyxl")
            return self._parse_xlsx(file_path)
        elif ext == '.xls':
            if not XLS_AVAILABLE:
                raise ImportError("XLS parsing requires xlrd: pip install xlrd")
            return self._parse_xls(file_path)
        else:
            raise ValueError(f"Unsupported Excel format: {ext}")

    def _parse_xlsx(self, file_path: Path) -> tuple[str, dict]:
        """Parse XLSX files using openpyxl."""
        workbook = load_workbook(str(file_path), read_only=True, data_only=True)

        sheets_content = []
        total_rows = 0

        for sheet_name in workbook.sheetnames:
            worksheet = workbook[sheet_name]
            sheet_content = [f"=== Sheet: {sheet_name} ==="]
            sheet_rows = 0

            for row in worksheet.iter_rows(values_only=True):
                if any(cell is not None for cell in row):
                    row_text = ' | '.join(str(cell) if cell is not None else '' for cell in row)
                    if row_text.strip():
                        sheet_content.append(row_text)
                        sheet_rows += 1

            if sheet_rows > 0:
                sheets_content.extend(sheet_content + [''])
                total_rows += sheet_rows

        workbook.close()

        return '\n'.join(sheets_content), {
            'sheets': len(workbook.sheetnames),
            'total_rows': total_rows
        }

    def _parse_xls(self, file_path: Path) -> tuple[str, dict]:
        """Parse XLS files using xlrd."""
        workbook = xlrd.open_workbook(str(file_path))

        sheets_content = []
        total_rows = 0

        for sheet_idx in range(workbook.nsheets):
            worksheet = workbook.sheet_by_index(sheet_idx)
            sheet_content = [f"=== Sheet: {worksheet.name} ==="]

            for row_idx in range(worksheet.nrows):
                row = worksheet.row_values(row_idx)
                row_text = ' | '.join(str(cell) for cell in row if cell)
                if row_text.strip():
                    sheet_content.append(row_text)

            if worksheet.nrows > 0:
                sheets_content.extend(sheet_content + [''])
                total_rows += worksheet.nrows

        return '\n'.join(sheets_content), {
            'sheets': workbook.nsheets,
            'total_rows': total_rows
        }

    def _parse_csv(self, file_path: Path) -> tuple[str, dict]:
        """Parse CSV files with automatic delimiter detection."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            # Detect dialect
            sample = f.read(8192)
            f.seek(0)

            try:
                dialect = csv.Sniffer().sniff(sample)
                reader = csv.reader(f, dialect)
            except csv.Error:
                f.seek(0)
                reader = csv.reader(f)

            rows = []
            for row_num, row in enumerate(reader):
                if row_num == 0:
                    rows.append(' | '.join(f"**{cell}**" for cell in row))
                else:
                    rows.append(' | '.join(row))

                if row_num > 10000:
                    rows.append(f"... (truncated after {row_num} rows)")
                    break

            return '\n'.join(rows), {'rows': len(rows) - 1}

    def _parse_tsv(self, file_path: Path) -> tuple[str, dict]:
        """Parse TSV files."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            reader = csv.reader(f, delimiter='\t')

            rows = []
            for row_num, row in enumerate(reader):
                if row_num == 0:
                    rows.append(' | '.join(f"**{cell}**" for cell in row))
                else:
                    rows.append(' | '.join(row))

                if row_num > 10000:
                    rows.append(f"... (truncated after {row_num} rows)")
                    break

            return '\n'.join(rows), {'rows': len(rows) - 1}

    def _parse_json(self, file_path: Path) -> tuple[str, dict]:
        """Parse JSON and JSONL files."""
        ext = file_path.suffix.lower()

        if ext in {'.jsonl', '.ndjson'}:
            # Line-delimited JSON
            content_lines = []
            with open(file_path, 'r', encoding='utf-8') as f:
                for line_num, line in enumerate(f):
                    line = line.strip()
                    if line:
                        try:
                            obj = json.loads(line)
                            content_lines.append(json.dumps(obj, indent=2))
                        except json.JSONDecodeError:
                            content_lines.append(f"Invalid JSON on line {line_num + 1}")

                    if line_num > 1000:
                        content_lines.append(f"... (truncated after {line_num} lines)")
                        break

            return '\n---\n'.join(content_lines), {'lines': len(content_lines)}
        else:
            # Regular JSON
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)

            content = json.dumps(data, indent=2, ensure_ascii=False)
            return content, {'json_type': type(data).__name__}

    def _parse_notebook(self, file_path: Path) -> tuple[str, dict]:
        """Parse Jupyter notebook files."""
        with open(file_path, 'r', encoding='utf-8') as f:
            notebook = json.load(f)

        content_parts = []
        code_cells = 0
        markdown_cells = 0

        for cell in notebook.get('cells', []):
            cell_type = cell.get('cell_type', 'unknown')
            source = cell.get('source', [])

            if isinstance(source, list):
                cell_content = ''.join(source)
            else:
                cell_content = str(source)

            if cell_content.strip():
                content_parts.append(f"=== {cell_type.title()} Cell ===")
                content_parts.append(cell_content)
                content_parts.append('')

                if cell_type == 'code':
                    code_cells += 1
                elif cell_type == 'markdown':
                    markdown_cells += 1

        return '\n'.join(content_parts), {
            'total_cells': code_cells + markdown_cells,
            'code_cells': code_cells,
            'markdown_cells': markdown_cells,
            'kernel': notebook.get('metadata', {}).get('kernelspec', {}).get('name', 'unknown')
        }

    def _parse_html(self, file_path: Path) -> tuple[str, dict]:
        """Parse HTML files using BeautifulSoup."""
        if not HTML_AVAILABLE:
            return self._parse_html_fallback(file_path)

        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()

        soup = BeautifulSoup(content, 'html.parser')

        # Remove script and style
        for script in soup(["script", "style"]):
            script.decompose()

        # Extract metadata
        title = soup.find('title')
        title_text = title.get_text() if title else ''

        meta_desc = soup.find('meta', attrs={'name': 'description'})
        meta_desc_text = meta_desc.get('content', '') if meta_desc else ''

        # Extract content
        main_content = soup.find('main') or soup.find('article') or soup.find('div', class_='content')
        text = main_content.get_text() if main_content else soup.get_text()

        # Clean
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)

        # Combine
        content_parts = []
        if title_text:
            content_parts.append(f"Title: {title_text}")
        if meta_desc_text:
            content_parts.append(f"Description: {meta_desc_text}")
        if text:
            content_parts.append(text)

        final_content = '\n\n'.join(content_parts)

        return final_content, {
            'title': title_text,
            'description': meta_desc_text,
            'links': len(soup.find_all('a')),
            'images': len(soup.find_all('img'))
        }

    def _parse_html_fallback(self, file_path: Path) -> tuple[str, dict]:
        """Fallback HTML parsing using regex."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()

        # Remove script and style
        content = re.sub(r'<script[^>]*>.*?</script>', '', content, flags=re.DOTALL | re.IGNORECASE)
        content = re.sub(r'<style[^>]*>.*?</style>', '', content, flags=re.DOTALL | re.IGNORECASE)

        # Extract title
        title_match = re.search(r'<title[^>]*>(.*?)</title>', content, re.IGNORECASE | re.DOTALL)
        title = title_match.group(1) if title_match else ''

        # Remove tags
        text = re.sub(r'<[^>]+>', ' ', content)
        text = re.sub(r'\s+', ' ', text).strip()

        final_content = f"Title: {title}\n\n{text}" if title else text

        return final_content, {'title': title, 'fallback_parser': True}

    def _parse_xml(self, file_path: Path) -> tuple[str, dict]:
        """Parse XML files including RSS and Atom feeds."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()

        if HTML_AVAILABLE:
            soup = BeautifulSoup(content, 'xml')
            text = soup.get_text()

            # Special handling for feeds
            if soup.find('rss') or soup.find('feed'):
                items = soup.find_all('item') or soup.find_all('entry')
                feed_content = []
                for item in items[:20]:
                    title = item.find('title')
                    description = item.find('description') or item.find('summary')
                    if title:
                        feed_content.append(f"Title: {title.get_text()}")
                    if description:
                        feed_content.append(f"Content: {description.get_text()}")
                    feed_content.append("---")

                text = '\n'.join(feed_content) if feed_content else text
        else:
            # Fallback
            text = re.sub(r'<[^>]+>', ' ', content)
            text = re.sub(r'\s+', ' ', text).strip()

        return text, {
            'xml_type': 'rss/atom' if 'rss' in content.lower() or 'feed' in content.lower() else 'xml'
        }

    def _parse_email(self, file_path: Path) -> tuple[str, dict]:
        """Parse email files (.eml, .msg)."""
        if not EMAIL_AVAILABLE:
            raise ImportError("Email parsing requires email library (standard library)")

        with open(file_path, 'rb') as f:
            parser = BytesParser(policy=email_default_policy)
            msg = parser.parse(f)

        # Extract components
        subject = msg.get('Subject', '')
        sender = msg.get('From', '')
        recipients = msg.get('To', '')
        date = msg.get('Date', '')

        # Extract body
        body = ''
        if msg.is_multipart():
            for part in msg.iter_parts():
                if part.get_content_type() == 'text/plain':
                    body += part.get_content()
                elif part.get_content_type() == 'text/html' and not body:
                    html_content = part.get_content()
                    if HTML_AVAILABLE:
                        soup = BeautifulSoup(html_content, 'html.parser')
                        body = soup.get_text()
                    else:
                        body = re.sub(r'<[^>]+>', ' ', html_content)
        else:
            body = msg.get_content()

        # Combine
        email_parts = []
        if subject:
            email_parts.append(f"Subject: {subject}")
        if sender:
            email_parts.append(f"From: {sender}")
        if recipients:
            email_parts.append(f"To: {recipients}")
        if date:
            email_parts.append(f"Date: {date}")
        if body:
            email_parts.append(f"\nContent:\n{body}")

        content = '\n'.join(email_parts)

        return content, {
            'subject': subject,
            'sender': sender,
            'recipients': recipients,
            'date': date
        }

    def _parse_rtf(self, file_path: Path) -> tuple[str, dict]:
        """Parse RTF files using pandoc or basic parsing."""
        if PANDOC_AVAILABLE:
            try:
                content = pypandoc.convert_file(str(file_path), 'plain', format='rtf')
                return content, {'parser': 'pandoc'}
            except Exception as e:
                logger.warning(f"Pandoc RTF parsing failed: {e}")

        # Basic fallback
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()

        # Basic RTF tag removal
        content = re.sub(r'\\[a-z]+\d*\s*', '', content)
        content = re.sub(r'[{}]', '', content)
        content = re.sub(r'\s+', ' ', content).strip()

        return content, {'parser': 'basic_rtf'}

    def _parse_archive(self, file_path: Path) -> tuple[str, dict]:
        """Parse archive files and extract text from contained files."""
        if not ARCHIVE_AVAILABLE:
            raise ImportError("Archive parsing requires zipfile and tarfile (standard library)")

        ext = file_path.suffix.lower()
        extracted_content = []
        file_list = []

        if ext == '.zip':
            with zipfile.ZipFile(file_path, 'r') as archive:
                for file_info in archive.filelist:
                    if not file_info.is_dir():
                        file_list.append(file_info.filename)
                        if file_info.file_size < 1024 * 1024:  # 1MB limit
                            try:
                                with archive.open(file_info.filename) as f:
                                    content = f.read(8192).decode('utf-8', errors='ignore')
                                    if content.strip():
                                        extracted_content.append(f"=== {file_info.filename} ===")
                                        extracted_content.append(content)
                            except:
                                continue

        elif ext in {'.tar', '.tar.gz', '.tgz', '.tar.bz2', '.tbz2', '.tar.xz', '.txz'}:
            with tarfile.open(file_path, 'r:*') as archive:
                for member in archive.getmembers():
                    if member.isfile():
                        file_list.append(member.name)
                        if member.size < 1024 * 1024:
                            try:
                                f = archive.extractfile(member)
                                if f:
                                    content = f.read(8192).decode('utf-8', errors='ignore')
                                    if content.strip():
                                        extracted_content.append(f"=== {member.name} ===")
                                        extracted_content.append(content)
                            except:
                                continue

        # Create summary
        summary_parts = [f"Archive contains {len(file_list)} files:"]
        summary_parts.extend([f"  - {name}" for name in file_list[:20]])
        if len(file_list) > 20:
            summary_parts.append(f"  ... and {len(file_list) - 20} more files")

        if extracted_content:
            summary_parts.extend(["", "=== Extracted Text Content ==="] + extracted_content)

        content = '\n'.join(summary_parts)

        return content, {
            'total_files': len(file_list),
            'extracted_files': len(extracted_content) // 2,
            'archive_type': ext
        }

    def _parse_text(self, file_path: Path) -> tuple[str, dict]:
        """Parse text and code files with encoding detection."""
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1']

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                    content = f.read()

                content = self._clean_text(content)

                return content, {
                    'encoding': encoding,
                    'lines': content.count('\n') + 1 if content else 0,
                    'characters': len(content)
                }

            except UnicodeDecodeError:
                continue

        raise ValueError("Could not decode file with any supported encoding")

    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        if not text:
            return ''

        # Remove excessive whitespace (but preserve some structure)
        text = re.sub(r'[ \t]+', ' ', text)

        # Remove control characters except newlines and tabs
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)

        # Normalize line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')

        # Remove excessive blank lines
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)

        return text.strip()


# ============================================================================
# Functional Interface (Convenience Functions)
# ============================================================================

# Global parser instance
_parser = FileParser()


def parse_file(file_path: Union[str, Path]) -> ParseResult:
    """
    Parse any supported file and extract text content.

    Args:
        file_path: Path to file to parse

    Returns:
        ParseResult with content, metadata, and status

    Example:
        >>> from shared.utils import parse_file
        >>> result = parse_file("document.pdf")
        >>> if result.success:
        ...     print(result.content[:200])
    """
    return _parser.parse_file(file_path)


def get_supported_extensions() -> Set[str]:
    """
    Get all supported file extensions.

    Returns:
        Set of supported extensions (e.g., {'.pdf', '.docx', ...})
    """
    return FileParser.get_supported_extensions()


def is_supported_file(file_path: Union[str, Path]) -> bool:
    """
    Check if file type is supported.

    Args:
        file_path: Path to file

    Returns:
        True if file type is supported

    Example:
        >>> from shared.utils import is_supported_file
        >>> if is_supported_file("example.docx"):
        ...     result = parse_file("example.docx")
    """
    return FileParser.is_supported(file_path)


def get_file_type(file_path: Union[str, Path]) -> str:
    """
    Get file type category.

    Args:
        file_path: Path to file

    Returns:
        File type category (text, code, document, etc.)

    Example:
        >>> from shared.utils import get_file_type
        >>> print(get_file_type("data.csv"))  # "data"
    """
    return FileParser.get_file_type(file_path)


# ============================================================================
# Testing
# ============================================================================

def _test_document_parsers():
    """Test function for standalone testing."""
    print("Testing Document Parsers...")
    print(f"\nSupported extensions: {len(get_supported_extensions())}")

    # Test categories
    test_files = {
        'text': Path('/tmp/test.txt'),
        'json': Path('/tmp/test.json'),
        'csv': Path('/tmp/test.csv'),
    }

    # Create test files
    test_files['text'].write_text("This is a test file.\nLine 2.\nLine 3.")
    test_files['json'].write_text('{"key": "value", "number": 42}')
    test_files['csv'].write_text("Name,Age,City\nAlice,30,NYC\nBob,25,LA")

    print("\n1. Testing file type detection...")
    for category, path in test_files.items():
        file_type = get_file_type(path)
        supported = is_supported_file(path)
        print(f"   {path.name}: type={file_type}, supported={supported}")

    print("\n2. Testing file parsing...")
    for category, path in test_files.items():
        result = parse_file(path)
        if result.success:
            print(f"   ✓ {path.name}: {len(result.content)} chars")
            print(f"     Metadata: {result.metadata}")
        else:
            print(f"   ✗ {path.name}: {result.error}")

    # Cleanup
    for path in test_files.values():
        path.unlink()

    print("\nAll tests complete!")


if __name__ == "__main__":
    _test_document_parsers()
